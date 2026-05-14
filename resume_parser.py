"""
resume_parser.py — ATS Analyzer Resume Parser
Handles PDF and DOCX text extraction entirely in-memory,
then runs SpaCy NER + regex to extract structured fields.
"""

from __future__ import annotations

import io
import re
import logging
from dataclasses import dataclass, field
from typing import Optional

# PDF extraction
from pdfminer.high_level import extract_text_to_fp
from pdfminer.layout import LAParams

# DOCX extraction
import docx

# NLP
import spacy
from spacy.language import Language

# Local
from security import sanitize_text_input, extract_pii_inventory

logger = logging.getLogger(__name__)

# ── SpaCy model (loaded once at module level) ─────────────────────────────────
try:
    NLP: Language = spacy.load("en_core_web_sm")
    logger.info("SpaCy model 'en_core_web_sm' loaded successfully.")
except OSError:
    NLP = None  # Handled gracefully in extraction calls
    logger.error(
        "SpaCy model 'en_core_web_sm' not found. "
        "Run: python -m spacy download en_core_web_sm"
    )


# ── Regex Patterns ────────────────────────────────────────────────────────────
_RE_EMAIL = re.compile(
    r"\b[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+\b"
)
_RE_PHONE = re.compile(
    r"(\+?\d{1,3}[\s\-.]?)?\(?\d{2,4}\)?[\s\-.]?\d{3,4}[\s\-.]?\d{4}"
)
_RE_LINKEDIN = re.compile(
    r"(https?://)?(www\.)?linkedin\.com/(in|pub|company)/[a-zA-Z0-9\-_%]+/?",
    re.IGNORECASE,
)
_RE_GITHUB = re.compile(
    r"(https?://)?(www\.)?github\.com/[a-zA-Z0-9\-_]+/?",
    re.IGNORECASE,
)
_RE_URL = re.compile(
    r"https?://[^\s<>\"{}|\\^`\[\]]+",
    re.IGNORECASE,
)

# Common section headers (used to identify document structure)
_SECTION_HEADERS = re.compile(
    r"^(experience|education|skills|projects|certifications|awards|"
    r"publications|summary|objective|profile|languages|interests|"
    r"volunteering|achievements|work history|professional experience)\s*:?\s*$",
    re.IGNORECASE | re.MULTILINE,
)

# ATS-hostile formatting indicators
_TABLE_INDICATOR = re.compile(
    r"(\|.+\|)|(\+[-+]+\+)",  # Markdown / ASCII tables
    re.MULTILINE,
)
_COLUMN_INDICATOR = re.compile(
    r"[ \t]{4,}",  # Multiple large gaps suggest multi-column layout
)


# ── Data Model ────────────────────────────────────────────────────────────────
@dataclass
class ParsedResume:
    """Structured output from the resume parser."""

    raw_text: str = ""
    clean_text: str = ""

    # Contact info (PII — stored for display to the user only)
    emails: list[str] = field(default_factory=list)
    phones: list[str] = field(default_factory=list)
    linkedin_urls: list[str] = field(default_factory=list)
    github_urls: list[str] = field(default_factory=list)
    other_urls: list[str] = field(default_factory=list)

    # NER extractions
    name: Optional[str] = None
    organizations: list[str] = field(default_factory=list)
    locations: list[str] = field(default_factory=list)
    skills: list[str] = field(default_factory=list)
    education_mentions: list[str] = field(default_factory=list)

    # Structural metadata
    sections_found: list[str] = field(default_factory=list)
    word_count: int = 0
    char_count: int = 0

    # ATS formatting flags
    has_tables: bool = False
    has_multi_columns: bool = False
    has_headers_footers: bool = False
    is_image_based: bool = False   # True when almost no text could be extracted

    # PII inventory (category → count, no actual values)
    pii_inventory: dict = field(default_factory=dict)

    # Parse error (None = success)
    error: Optional[str] = None


# ── Extraction ────────────────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> ParsedResume:
    """
    Top-level entry point.  Dispatches to PDF or DOCX parser based on
    the file extension, then runs NLP enrichment.

    Args:
        file_bytes:  Raw bytes of the uploaded file (from in-memory buffer).
        filename:    Original filename (used to determine parser).

    Returns:
        A fully populated ParsedResume dataclass.
    """
    result = ParsedResume()
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if ext == "pdf":
            raw_text = _extract_pdf(file_bytes)
        elif ext == "docx":
            raw_text = _extract_docx(file_bytes)
        else:
            result.error = f"Unsupported extension: .{ext}"
            return result
    except Exception as exc:
        logger.exception("Text extraction failed for '%s'", filename)
        result.error = f"Could not parse file: {exc}"
        return result

    if not raw_text or len(raw_text.strip()) < 50:
        result.is_image_based = True
        result.error = (
            "Very little text was extracted. The resume may be image-based "
            "(scanned PDF). Please convert to a text-based PDF or DOCX."
        )
        return result

    # Sanitize for safe downstream processing
    clean = sanitize_text_input(raw_text)

    result.raw_text = raw_text
    result.clean_text = clean
    result.word_count = len(clean.split())
    result.char_count = len(clean)
    result.pii_inventory = extract_pii_inventory(clean)

    # Regex extractions
    _extract_contact_info(clean, result)

    # Structural analysis
    _detect_formatting_issues(raw_text, result)
    _detect_sections(clean, result)

    # NLP enrichment
    _run_ner(clean, result)

    return result


# ── PDF Extraction ────────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF using pdfminer.six.
    Entirely in-memory — no temporary files written.
    """
    buf_in = io.BytesIO(file_bytes)
    buf_out = io.StringIO()
    laparams = LAParams(
        line_overlap=0.5,
        char_margin=2.0,
        line_margin=0.5,
        word_margin=0.1,
        detect_vertical=False,
    )
    extract_text_to_fp(buf_in, buf_out, laparams=laparams, output_type="text")
    text = buf_out.getvalue()
    logger.info("PDF extracted: %d chars", len(text))
    return text


# ── DOCX Extraction ───────────────────────────────────────────────────────────

def _extract_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Reads from an in-memory BytesIO buffer.
    """
    buf = io.BytesIO(file_bytes)
    doc = docx.Document(buf)
    parts: list[str] = []

    # Paragraphs (main body)
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)

    # Tables (extract cell text so NLP can use it even though ATS may struggle)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    parts.append(txt)

    text = "\n".join(parts)
    logger.info("DOCX extracted: %d chars", len(text))
    return text


# ── Regex Contact Extraction ──────────────────────────────────────────────────

def _extract_contact_info(text: str, result: ParsedResume) -> None:
    result.emails = list({m.group() for m in _RE_EMAIL.finditer(text)})

    # Phones: deduplicate by normalising digits
    raw_phones = [m.group().strip() for m in _RE_PHONE.finditer(text)]
    seen: set[str] = set()
    for p in raw_phones:
        key = re.sub(r"\D", "", p)
        if len(key) >= 7 and key not in seen:
            seen.add(key)
            result.phones.append(p)

    result.linkedin_urls = list(
        {m.group() for m in _RE_LINKEDIN.finditer(text)}
    )
    result.github_urls = list(
        {m.group() for m in _RE_GITHUB.finditer(text)}
    )

    # Other URLs (exclude already captured linkedin/github)
    all_urls = {m.group() for m in _RE_URL.finditer(text)}
    result.other_urls = [
        u for u in all_urls
        if not _RE_LINKEDIN.search(u) and not _RE_GITHUB.search(u)
    ]


# ── Formatting / ATS Issue Detection ─────────────────────────────────────────

def _detect_formatting_issues(raw_text: str, result: ParsedResume) -> None:
    result.has_tables = bool(_TABLE_INDICATOR.search(raw_text))

    # Multi-column: many lines with 4+ spaces gap hints at side-by-side layout
    column_lines = sum(
        1 for line in raw_text.splitlines()
        if _COLUMN_INDICATOR.search(line)
    )
    result.has_multi_columns = column_lines > (len(raw_text.splitlines()) * 0.15)

    # Headers/footers heuristic: repeated short lines at start/end of pages
    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
    if lines:
        result.has_headers_footers = (
            len(lines[0]) < 60 and len(lines[-1]) < 60
        )


def _detect_sections(text: str, result: ParsedResume) -> None:
    matches = _SECTION_HEADERS.findall(text)
    result.sections_found = list({m[0].title() for m in matches if m[0]})


# ── NLP (SpaCy NER) ───────────────────────────────────────────────────────────

# Education keywords for heuristic detection when NER misses them
_EDUCATION_KW = re.compile(
    r"\b(bachelor|master|phd|doctorate|b\.?sc|m\.?sc|mba|b\.?tech|m\.?tech"
    r"|b\.?e\b|m\.?e\b|diploma|degree|university|college|institute|school"
    r"|12th|10th|high school|secondary|undergraduate|graduate|postgraduate"
    r"|bca|mca|bba|b\.?com|m\.?com)\b",
    re.IGNORECASE,
)

# Common technical skill terms (supplement NER)
_SKILL_KW = re.compile(
    r"\b(python|java|javascript|typescript|c\+\+|c#|ruby|go|rust|swift|kotlin"
    r"|sql|nosql|mongodb|postgresql|mysql|redis|elasticsearch"
    r"|aws|azure|gcp|docker|kubernetes|terraform|ansible|jenkins|ci/cd"
    r"|react|angular|vue|node\.?js|django|flask|fastapi|spring|\.net"
    r"|machine learning|deep learning|nlp|computer vision|tensorflow|pytorch"
    r"|pandas|numpy|scikit.?learn|spark|hadoop|kafka|airflow"
    r"|git|github|gitlab|jira|agile|scrum|rest|graphql|grpc"
    r"|linux|bash|powershell|html|css|sass"
    r"|communication|leadership|teamwork|problem.solving|critical thinking"
    r"|project management|data analysis|research|presentation)\b",
    re.IGNORECASE,
)


def _run_ner(text: str, result: ParsedResume) -> None:
    """Run SpaCy NER and heuristic keyword extraction."""
    if NLP is None:
        logger.warning("SpaCy model unavailable; skipping NER.")
        # Fall back to pure regex skills/education
        result.skills = list(
            {m.group().lower() for m in _SKILL_KW.finditer(text)}
        )
        result.education_mentions = list(
            {m.group().lower() for m in _EDUCATION_KW.finditer(text)}
        )
        return

    # SpaCy has a 1M char limit by default; chunk if necessary
    chunk = text[:950_000]
    doc = NLP(chunk)

    orgs, locs, persons = set(), set(), set()
    for ent in doc.ents:
        label = ent.label_
        val = ent.text.strip()
        if not val or len(val) < 2:
            continue
        if label == "ORG":
            orgs.add(val)
        elif label in ("GPE", "LOC"):
            locs.add(val)
        elif label == "PERSON" and not result.name:
            persons.add(val)

    result.organizations = sorted(orgs)
    result.locations = sorted(locs)

    # Heuristic: take the first PERSON entity as the candidate name
    if persons:
        result.name = sorted(persons, key=lambda x: text.find(x))[0]

    # Skills: regex keyword matching (NER doesn't reliably label skills)
    result.skills = sorted(
        {m.group().lower() for m in _SKILL_KW.finditer(text)}
    )

    # Education: regex for degree/institution keywords
    result.education_mentions = sorted(
        {m.group().lower() for m in _EDUCATION_KW.finditer(text)}
    )
